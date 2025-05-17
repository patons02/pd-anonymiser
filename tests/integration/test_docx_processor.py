import os

import pytest

from docx.api import Document as DocxDocument
from pd_anonymiser.docx_processor import anonymise_docx_comprehensive
from pd_anonymiser.anonymiser import AnonymisationResult
from pd_anonymiser.xml_helpers import extract_comments, extract_footnotes, extract_textboxes

@pytest.fixture
def fixture_paths():
    base = os.path.dirname(__file__)
    fixtures_dir = os.path.join(base, 'fixtures')
    input_docx = os.path.join(fixtures_dir, 'comments_footnotes_textboxes.docx')
    expected_txt = os.path.join(fixtures_dir, 'expected_output.txt')
    return input_docx, expected_txt

def test_docx_anonymisation_integration(tmp_path, fixture_paths):
    input_docx, expected_txt = fixture_paths
    output_path = tmp_path / 'anonymised.docx'

    # Run the full anonymisation
    result = anonymise_docx_comprehensive(
        input_path=input_docx,
        output_path=str(output_path),
    )

    # Check result metadata
    assert isinstance(result, AnonymisationResult)
    assert result.session_id is not None
    assert result.key is not None

    # Load expected lines
    with open(expected_txt, 'r') as f:
        expected_lines = [line.strip() for line in f if line.strip()]

    # Collect actual lines in the same order as the processor
    doc = DocxDocument(str(output_path))
    actual_lines = []

    def collect_paras(paragraphs):
        for p in paragraphs:
            txt = p.text.strip()
            if txt:
                actual_lines.append(txt)

    # Body paragraphs
    collect_paras(doc.paragraphs)
    # Tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                collect_paras(cell.paragraphs)
    # Headers and footers
    for sec in doc.sections:
        collect_paras(sec.header.paragraphs)
        collect_paras(sec.footer.paragraphs)
    # Comments, footnotes, textboxes
    comments_texts, _ = extract_comments(doc.part.comments_part)
    actual_lines.extend(comments_texts)
    footnotes_texts, _ = extract_footnotes(doc.part.footnotes_part)
    actual_lines.extend(footnotes_texts)
    textbox_texts, _ = extract_textboxes(doc.part.element)
    actual_lines.extend(textbox_texts)

    assert actual_lines == expected_lines

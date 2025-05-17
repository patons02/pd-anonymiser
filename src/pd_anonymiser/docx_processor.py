from docx import Document
from pd_anonymiser.anonymiser import anonymise_text, AnonymisationResult
from pd_anonymiser.xml_helpers import (
    extract_comments,
    extract_footnotes,
    extract_textboxes,
)
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def _anonymise_and_replace(paragraphs, result_iter):
    for para in paragraphs:
        if para.text.strip():
            para.text = next(result_iter)


def anonymise_docx_comprehensive(
    input_path: str,
    output_path: str,
    language: str = "en",
    use_reusable_tags: bool = True,
    model: str = "all",
    allow_reidentification: bool = False,
) -> AnonymisationResult:
    doc = Document(input_path)

    # Collect all text: paragraphs, tables, headers, footers, XML parts
    text_blocks = []

    def collect_text_from_paragraphs(paragraphs):
        for para in paragraphs:
            if para.text.strip():
                text_blocks.append(para.text)

    # Body
    collect_text_from_paragraphs(doc.paragraphs)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                collect_text_from_paragraphs(cell.paragraphs)

    # Headers and footers
    for section in doc.sections:
        collect_text_from_paragraphs(section.header.paragraphs)
        collect_text_from_paragraphs(section.footer.paragraphs)

    # Comments & Footnotes come in via OPC relationships
  # --- Comments ---
    comment_nodes = []

    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            comments_text, comment_nodes = extract_comments(rel.target_part)
            text_blocks.extend(comments_text)
            break

    # --- Footnotes ---
    footnote_nodes = []
    for rel in doc.part.rels.values():
        if rel.reltype == RT.FOOTNOTES:
            footnotes_text, footnote_nodes = extract_footnotes(rel.target_part)
            text_blocks.extend(footnotes_text)
            break

    # --- Textboxes (still via the main document XML) ---
    textbox_text, textbox_nodes = extract_textboxes(doc.part.element)
    text_blocks.extend(textbox_text)

    # Anonymise all text at once
    full_text = "\n".join(text_blocks)
    result = anonymise_text(
        text=full_text,
        language=language,
        use_reusable_tags=use_reusable_tags,
        model=model,
        allow_reidentification=allow_reidentification,
    )

    # Replace body
    anonymized_lines = result.text.split("\n")
    result_iter = iter(anonymized_lines)

    _anonymise_and_replace(doc.paragraphs, result_iter)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _anonymise_and_replace(cell.paragraphs, result_iter)

    for section in doc.sections:
        _anonymise_and_replace(section.header.paragraphs, result_iter)
        _anonymise_and_replace(section.footer.paragraphs, result_iter)

    # Replace XML nodes
    for node in comment_nodes:
        node.text = next(result_iter)
    for node in footnote_nodes:
        node.text = next(result_iter)
    for node in textbox_nodes:
        node.text = next(result_iter)

    doc.save(output_path)
    return result

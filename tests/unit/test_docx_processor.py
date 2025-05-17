import pytest

# Import the real python-docx Document
from docx.api import Document as DocxDocument

import pd_anonymiser.docx_processor as docx_processor
from pd_anonymiser.anonymiser import AnonymisationResult


class FakeNode:
    def __init__(self):
        self.text = None


@pytest.fixture(autouse=True)
def _patch_document(monkeypatch):
    """
    Whenever docx_processor.Document(...) is called inside the code under test,
    substitute it with a fake that delegates to the real DocxDocument and
    then injects empty comments_part & footnotes_part so the XML branches run.
    """

    def fake_document(path):
        # load the real file
        doc = DocxDocument(path)

        # give it dummy parts so extract_comments/extract_footnotes don't skip
        class DummyPart:
            pass

        doc.part.comments_part = DummyPart()
        doc.part.footnotes_part = DummyPart()
        return doc

    # patch into the module under test
    monkeypatch.setattr(docx_processor, "Document", fake_document)


@pytest.fixture
def sample_docx(tmp_path):
    """Create a simple .docx with a body paragraph, table, header and footer."""
    path = tmp_path / "input.docx"
    doc = DocxDocument()
    # Body
    doc.add_paragraph("BodyText")
    # Table
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "TableText"
    # Header
    hdr = doc.sections[0].header
    hdr.add_paragraph("HeaderText")
    # Footer
    ftr = doc.sections[0].footer
    ftr.add_paragraph("FooterText")
    doc.save(str(path))
    return str(path)


def test_anonymise_docx_comprehensive_basic(monkeypatch, tmp_path, sample_docx):
    # fake XML nodes
    fake_comment_node = FakeNode()
    fake_footnote_node = FakeNode()
    fake_textbox_node = FakeNode()

    # patch XML extractor calls
    monkeypatch.setattr(
        docx_processor,
        "extract_comments",
        lambda part: (["CommentText"], [fake_comment_node]),
    )
    monkeypatch.setattr(
        docx_processor,
        "extract_footnotes",
        lambda part: (["FootnoteText"], [fake_footnote_node]),
    )
    monkeypatch.setattr(
        docx_processor,
        "extract_textboxes",
        lambda part: (["TextboxText"], [fake_textbox_node]),
    )

    # prepare fake anonymisation result
    expected = ["A1", "B1", "C1", "D1", "E1", "F1", "G1"]
    fake_result = AnonymisationResult(
        text="\n".join(expected),
        session_id="fake-session",
        key="fake-key",
    )
    monkeypatch.setattr(docx_processor, "anonymise_text", lambda **kwargs: fake_result)

    out_path = tmp_path / "output.docx"
    result = docx_processor.anonymise_docx_comprehensive(
        input_path=sample_docx,
        output_path=str(out_path),
    )

    # verify return value
    assert result is fake_result
    assert result.text == "\n".join(expected)
    assert result.session_id == "fake-session"
    assert result.key == "fake-key"

    # load the output and check replacements
    doc_out = DocxDocument(str(out_path))
    assert doc_out.paragraphs[0].text == "A1"
    assert doc_out.tables[0].cell(0, 0).text == "B1"

    header_non_empty = [
        p for p in doc_out.sections[0].header.paragraphs if p.text.strip()
    ]
    assert header_non_empty[-1].text == "C1"

    footer_non_empty = [
        p for p in doc_out.sections[0].footer.paragraphs if p.text.strip()
    ]
    assert footer_non_empty[-1].text == "D1"

    # XML nodes got their .text set
    assert fake_comment_node.text == "E1"
    assert fake_footnote_node.text == "F1"
    assert fake_textbox_node.text == "G1"


def test_anonymise_docx_comprehensive_no_xml_nodes(monkeypatch, tmp_path, sample_docx):
    # patch XML extractors to return nothing
    monkeypatch.setattr(docx_processor, "extract_comments", lambda part: ([], []))
    monkeypatch.setattr(docx_processor, "extract_footnotes", lambda part: ([], []))
    monkeypatch.setattr(docx_processor, "extract_textboxes", lambda part: ([], []))

    expected = ["X1", "X2", "X3", "X4"]
    fake_result = AnonymisationResult(
        text="\n".join(expected),
        session_id=None,
        key=None,
    )
    monkeypatch.setattr(docx_processor, "anonymise_text", lambda **kwargs: fake_result)

    out_path = tmp_path / "output_no_xml.docx"
    result = docx_processor.anonymise_docx_comprehensive(
        input_path=sample_docx,
        output_path=str(out_path),
    )

    # return value
    assert result is fake_result
    assert result.text == "\n".join(expected)
    assert result.session_id is None
    assert result.key is None

    # verify replacements
    doc_out = DocxDocument(str(out_path))
    assert doc_out.paragraphs[0].text == "X1"
    assert doc_out.tables[0].cell(0, 0).text == "X2"

    header_non_empty = [
        p for p in doc_out.sections[0].header.paragraphs if p.text.strip()
    ]
    assert header_non_empty[-1].text == "X3"

    footer_non_empty = [
        p for p in doc_out.sections[0].footer.paragraphs if p.text.strip()
    ]
    assert footer_non_empty[-1].text == "X4"
